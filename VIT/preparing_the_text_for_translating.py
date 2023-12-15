import docx


# bu fonksiyon asıl işi yapan fonksiyonumuz. girdi metnini bir docx dosyasından okunmaya, istenmeyen varlıklarını yok
# etmeye veya istediğimiz bir karakterle değiştirmeye yardım ediyor.
def text_correction(file_path):
    document = docx.Document(file_path)
    paragraphs = []
    text = ''

    for paragraph in document.paragraphs:
        paragraphs.append(paragraph.text)

    # paragrafları sırasıyla text stringine yerleştiriyorum. Bunu yaparken paragraflar arasına -\n yerine -
    # 'tek boşluk' karakterini de ekliyorum.
    for i in paragraphs:
        text += i.strip('\n') + ' '
    # text değişkenindeki yazılarda bulunan 'tab' karakterini 'bir boşluk' karakteriyle değiştiriyorum.
    text = text.replace('\t', ' ')

    # text değşkeni içinde bulunan arka arkaya gelmiş boşluk karakterlerini tespit edip, sayıyor ve bunları 'tek
    # boşluk' karakteriyle değiştiriyorum.
    space = 0
    for j in text:
        if j == ' ':
            space += 1
        # burada bir extra işlem kaybımız var. Yalnız, while içine girip işlemi gerçekleştirdikten sonra space
        # sayısını 0 yapmak 1 yapmaktan daha az işlem kaybına yol açıyor. bu sebeple 0 seçildi. İşlem gücü kaybı:
        # aslında iki kelime arasında tek boşluk olmasına rağmen space değişkenimiz her boşluk karakterinde arttığı
        # için bu iki boşlukta bir while döngüsüne girilerek değişiklik yapılmaya çalışılıyor!
        while space > 1:
            text = text.replace((space * ' '), ' ')
            space = 0
    # print(text)  # metin dosyasına yazılan yazı ekrana basılıyor. gereksiz!
    return text


# Düzeltilmiş metni '.txt' uzantılı dosyaya yazdırıyoruz.
def write2txt_file(file_path, text):
    with open(file_path, 'w') as text_output_file:
        text_output_file.writelines(text)
    return f'Düzeltme tamamlandı! {output_file_1} yolundaki çıkış dosyasını kontrol ediniz.'  # Burası kullanıcıya
    # bilgi vermek için. Aslında çok gerekli değil. Fonksiyonun return  değerini "True" ya da "1" olarak da
    # ayarlayabilirdik.


# Düzeltilmiş metni '.docx' uzantılı dosyaya yazdırıyoruz.
def write2docx(file_path, text):
    document = docx.Document()
    document.add_paragraph(text)
    document.save(file_path)
    return f'Düzeltme tamamlandı! {output_file_2} yolundaki çıkış dosyasını kontrol ediniz.'  # Burası kullanıcıya
    # bilgi vermek için. Aslında çok gerekli değil. Fonksiyonun return  değerini "True" ya da "1" olarak da
    # ayarlayabilirdik.


if __name__ == '__main__':
    input_file = 'C:\\Users\\user\\Desktop\\giris.docx'
    output_file_1 = 'C:\\Users\\user\\Desktop\\sonuc.txt'
    output_file_2 = 'C:\\Users\\user\\Desktop\\docx_file.docx'

    print(write2txt_file(output_file_1, text_correction(input_file)))
    print(write2docx(output_file_2, text_correction(input_file)))

