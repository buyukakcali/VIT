"""
Readme:
 Kodun Amacı; çeviri yapılmak istenen metnin uzun olması ve genellikle bir pdf, word veya başka metin tabanlı bir dosyadan çeviri yapılmak istenmesi durumunda karşılaşılan zorlukları azaltmak, çevrilen metnin daha hızlı ve kolay biçimde anlaşılmasını sağlamaya yardımcı olmaktır. Çünkü başka formatlarda saklı tutulan bu metinlerde aslında olmaması gereken 'paragraf', 'tab' gibi çeviri sayfalarında çeviri kalitesini etkileyen unsurlar bulunmaktadır. Kod, farklı dosyalardan kopyaladığımız metni bu unsurlardan arındırmakta ve yazıyı nesir haline getirmektedir. Böylece daha doğru bir çeviri yapılmakta ve kullanıcısı için zaman tasarrufu sağlamaktadır.

Açıklama:
 - text_correction fonksiyonu dosya yolu parametresini isteyen, bir word dosyasından metinleri alıp iki aşamada string değişkene aktaran ve stringdeki 'paragraf', 'tab' kayıtlarını 'bir boşluk' karakteri ile değiştiren, ayrıca birden fazla boşluk olması halinde bunu tek boşluğa dönüştürerek bu stringi döndüren bir fonksiyondur.
 - write2txt_file fonksiyonu dosya yolu ve text_correction fonksiyonundan elde ettiği stringi parametre olarak alıp ilgili yolda txt uzantılı bir metin belegesine bu stringi kaydetmektedir.
 - write2docx_file fonksiyonu dosya yolu ve text_correction fonksiyonundan elde ettiği stringi parametre olarak alıp ilgili yolda docx uzantılı bir Word belegesine bu stringi kaydetmektedir.

Parametreler:
 - text_correction fonksiyonu için: file_path
 -  write2txt_file fonksiyonu için: file_path ve text
 -  write2docx_file fonksiyonu için: file_path ve text

Kısıtlamalar:
 - text_correction: yalnızca word dosyasını okuyabilir.
 - Stringdeki paragrafları, girintileri ve metnin orijinal şeklini ortadan kaldırarak yazılanların kolaylıkla y-ayırt edilmesini zorlaştırır.
 - Okunacak dosyanın tam yolunun koda elle girilmesi gerekir. Dosya yöneticisi ile seçmek mümkün değildir.
 - Kodda hata kontrolü yapıları (try-catch) kullanılmamıştır. Kod kolaylıkla çökebilir. Mesela girdi dosya yolunda ilgili dosya bulunmazsa kod hata döndürür, ancak bu hatanın ne olduğu kod geliştiricisi tarafından açık bir şekilde belirtilmemiştir.

Yapılar:
 - for döngüsü çoğunlukla kullanılmıştır. 
 - while döngüsü yalnızca bir yerde kullanılmıştır.
 - docx modülü marifetiyle bir Document nesnesi oluşturularak dosya okunmuştur.
 - with open yapısı dosya yazdırma işlemi için iki ayrı fonksiyonda kullanılmıştır.

Çıktı:
 -düzeltme işlemi tamamlandıktan sonra string metin dosyası ve word dosyasına yazılmıştır.
 - Ekrana ıktıların yolu ile birlikte işlemin tamamlandığı bilgisi yazdırılmıştır.
 """


import docx


# bu fonksiyon asıl işi yapan fonksiyonumuz. girdi metnini bir docx dosyasından okunmaya, istenmeyen varlıklarını yok
# etmeye veya istediğimiz bir karakterle değiştirmeye yardım ediyor.
def text_correction(file_path):
    document = docx.Document(file_path)
    paragraphs = []
    text = ''

    for paragraph in document.paragraphs:
        paragraphs.append(paragraph.text)

    # paragrafları sırasıyla text stringine yerleştiriyorum. Bunu yaparken paragraflar arasına -\n yerine -
    # 'tek boşluk' karakterini de ekliyorum.
    for i in paragraphs:
        text += i.strip('\n') + ' '
    # text değişkenindeki yazılarda bulunan 'tab' karakterini 'bir boşluk' karakteriyle değiştiriyorum.
    text = text.replace('\t', ' ')

    # text değşkeni içinde bulunan arka arkaya gelmiş boşluk karakterlerini tespit edip, sayıyor ve bunları 'tek
    # boşluk' karakteriyle değiştiriyorum.
    space = 0
    for j in text:
        if j == ' ':
            space += 1
        # burada bir extra işlem kaybımız var. Yalnız, while içine girip işlemi gerçekleştirdikten sonra space
        # sayısını 0 yapmak 1 yapmaktan daha az işlem kaybına yol açıyor. bu sebeple 0 seçildi. İşlem gücü kaybı:
        # aslında iki kelime arasında tek boşluk olmasına rağmen space değişkenimiz her boşluk karakterinde arttığı
        # için bu iki boşlukta bir while döngüsüne girilerek değişiklik yapılmaya çalışılıyor!
        while space > 1:
            text = text.replace((space * ' '), ' ')
            space = 0
    # print(text)  # metin dosyasına yazılan yazı ekrana basılıyor. gereksiz!
    return text


# Düzeltilmiş metni '.txt' uzantılı dosyaya yazdırıyoruz.
def write2txt_file(file_path, text):
    with open(file_path, 'w') as text_output_file:
        text_output_file.writelines(text)
    return f'Düzeltme tamamlandı! {output_file_1} yolundaki çıkış dosyasını kontrol ediniz.'  # Burası kullanıcıya
    # bilgi vermek için. Aslında çok gerekli değil. Fonksiyonun return  değerini "True" ya da "1" olarak da
    # ayarlayabilirdik.


# Düzeltilmiş metni '.docx' uzantılı dosyaya yazdırıyoruz.
def write2docx_file(file_path, text):
    document = docx.Document()
    document.add_paragraph(text)
    document.save(file_path)
    return f'Düzeltme tamamlandı! {output_file_2} yolundaki çıkış dosyasını kontrol ediniz.'  # Burası kullanıcıya
    # bilgi vermek için. Aslında çok gerekli değil. Fonksiyonun return  değerini "True" ya da "1" olarak da
    # ayarlayabilirdik.


if __name__ == '__main__':
    input_file = 'C:\\Users\\user\\Desktop\\giris.docx'
    output_file_1 = 'C:\\Users\\user\\Desktop\\sonuc1.txt'
    output_file_2 = 'C:\\Users\\user\\Desktop\\sonuc2.docx'

    print(write2txt_file(output_file_1, text_correction(input_file)))
    print(write2docx_file(output_file_2, text_correction(input_file)))
